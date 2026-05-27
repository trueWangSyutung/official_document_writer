#!/usr/bin/env python3
"""
基于国标GB/T 9704-2012 的word公文排版

支持平台：macOS, Windows, Linux
用途：将word文档格式化为GB/T 9704-2012规范的格式

使用方法：
python format.py [input_file] [output_file]

参数说明：
input_file: 输入文件路径
output_file: 输出文件路径
"""
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING


class WordFormatter:
    """基于GB/T 9704-2012标准的Word文档格式化工具"""

    # GB/T 9704-2012 页面设置常量 (单位：厘米)
    PAGE_WIDTH = Cm(21)      # A4纸宽度
    PAGE_HEIGHT = Cm(29.7)   # A4纸高度
    MARGIN_TOP = Cm(3.7)     # 上边距
    MARGIN_BOTTOM = Cm(3.5)  # 下边距
    MARGIN_LEFT = Cm(2.8)    # 左边距
    MARGIN_RIGHT = Cm(2.6)   # 右边距

    # 字体设置
    BODY_FONT_NAME = "仿宋_GB2312"
    TITLE_FONT_NAME = "方正小标宋简体"
    RED_HEAD_FONT_NAME = "黑体"
    FONT_SIZE_BODY = Pt(16)     # 三号字
    FONT_SIZE_TITLE = Pt(22)    # 二号字
    FONT_SIZE_HEADING = Pt(16)  # 三号字
    FONT_SIZE_RED_HEAD = Pt(22) # 红头单位名称字号
    FONT_SIZE_RED_HEAD_NUM = Pt(16)  # 红头编号字号

    # 行间距
    LINE_SPACING = Pt(28)
    LINE_SPACING_RULE = WD_LINE_SPACING.EXACTLY

    # 红头颜色 (红色系)
    RED_HEAD_COLOR = "#C41E3A"

    def __init__(self, input_file: Path, output_file: Path, 
                 org_name="国央企", doc_title="", doc_number="", doc_date=""):
        self.input_file = input_file
        self.output_file = output_file
        self.org_name = org_name
        self.doc_title = doc_title
        self.doc_number = doc_number
        self.doc_date = doc_date
        self.doc = Document(self.input_file)

        if not self.input_file.exists():
            raise FileNotFoundError(f"输入文件不存在: {self.input_file}")

        self._setup_page_layout()
        self._setup_styles()

    def _setup_page_layout(self):
        """设置页面布局"""
        sections = self.doc.sections
        for section in sections:
            section.page_width = self.PAGE_WIDTH
            section.page_height = self.PAGE_HEIGHT
            section.top_margin = self.MARGIN_TOP
            section.bottom_margin = self.MARGIN_BOTTOM
            section.left_margin = self.MARGIN_LEFT
            section.right_margin = self.MARGIN_RIGHT

    def _setup_styles(self):
        """设置文档样式"""
        try:
            normal_style = self.doc.styles['Normal']
            normal_style.font.name = self.BODY_FONT_NAME
            normal_style.font.size = self.FONT_SIZE_BODY
        except Exception as e:
            print(f"样式设置警告: {e}")

    def format(self, add_red_head=False):
        """执行文档格式化"""
        if add_red_head:
            self._add_red_head()
        self._format_paragraphs()
        self._format_tables()
        self._save_document()

    def _add_red_head(self):
        """添加红头标识 - 符合GB/T 9704-2012规范"""
        if not self.org_name:
            return

        body = self.doc._body._body
        
        red_head_elements = []
        
        if self.doc_number:
            num_para = self.doc.add_paragraph()
            num_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            num_run = num_para.add_run(self.doc_number)
            num_run.font.name = self.BODY_FONT_NAME
            num_run.font.size = self.FONT_SIZE_RED_HEAD_NUM
            num_run.font.color.rgb = self._hex_to_rgb(self.RED_HEAD_COLOR)
            rPr = num_run._element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:eastAsia'), self.BODY_FONT_NAME)
            num_para.paragraph_format.space_after = Pt(5)
            red_head_elements.append(num_para._element)

        org_para = self.doc.add_paragraph()
        org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        org_run = org_para.add_run(self.org_name)
        org_run.font.name = self.RED_HEAD_FONT_NAME
        org_run.font.size = self.FONT_SIZE_RED_HEAD
        org_run.font.color.rgb = self._hex_to_rgb(self.RED_HEAD_COLOR)
        rPr = org_run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), self.RED_HEAD_FONT_NAME)
        org_para.paragraph_format.space_after = Pt(10)
        red_head_elements.append(org_para._element)
        
        if self.doc_title:
            title_para = self.doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run(self.doc_title)
            title_run.font.name = self.TITLE_FONT_NAME
            title_run.font.size = self.FONT_SIZE_TITLE
            title_run.font.color.rgb = self._hex_to_rgb(self.RED_HEAD_COLOR)
            rPr = title_run._element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:eastAsia'), self.TITLE_FONT_NAME)
            title_para.paragraph_format.space_after = Pt(25)
            red_head_elements.append(title_para._element)
        
        line_para = self.doc.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line_run = line_para.add_run("─" * 60)
        line_run.font.color.rgb = self._hex_to_rgb(self.RED_HEAD_COLOR)
        line_para.paragraph_format.space_after = Pt(15)
        red_head_elements.append(line_para._element)
        
        for element in reversed(red_head_elements):
            body.insert(0, element)

    def _hex_to_rgb(self, hex_color):
        """将十六进制颜色转换为RGB对象"""
        hex_color = hex_color.lstrip('#')
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)
        from docx.shared import RGBColor
        return RGBColor(r, g, b)

    def _format_paragraphs(self):
        """格式化所有段落"""
        for paragraph in self.doc.paragraphs:
            self._format_paragraph(paragraph)

    def _format_paragraph(self, paragraph):
        """格式化单个段落"""
        # 设置字体
        for run in paragraph.runs:
            run.font.name = self.BODY_FONT_NAME
            run.font.size = self.FONT_SIZE_BODY

            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:eastAsia'), self.BODY_FONT_NAME)

        # 设置行间距
        paragraph.paragraph_format.line_spacing = self.LINE_SPACING
        paragraph.paragraph_format.line_spacing_rule = self.LINE_SPACING_RULE

        # 首行缩进（除了标题和空行）
        text = paragraph.text.strip()
        if text and not self._is_title(paragraph):
            if not paragraph.paragraph_format.first_line_indent:
                paragraph.paragraph_format.first_line_indent = Cm(0.85)  # 2字符

        # 标题居中
        if self._is_title(paragraph):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = self.TITLE_FONT_NAME
                run.font.size = self.FONT_SIZE_TITLE

                rPr = run._element.get_or_add_rPr()
                rFonts = rPr.get_or_add_rFonts()
                rFonts.set(qn('w:eastAsia'), self.TITLE_FONT_NAME)

    def _is_title(self, paragraph) -> bool:
        """判断是否为标题段落"""
        if not paragraph.text.strip():
            return False

        text = paragraph.text.strip()
        if len(text) <= 20 and not text.endswith(('。', '，', '：', '；')):
            alignment = paragraph.alignment
            return alignment == WD_ALIGN_PARAGRAPH.CENTER

        return False

    def _format_tables(self):
        """格式化所有表格"""
        for table in self.doc.tables:
            self._format_table(table)

    def _format_table(self, table):
        """格式化单个表格"""
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self._format_paragraph(paragraph)

                for run in cell.paragraphs[0].runs if cell.paragraphs else []:
                    run.font.size = self.FONT_SIZE_BODY
                    run.font.name = self.BODY_FONT_NAME

    def _save_document(self):
        """保存文档"""
        self.doc.save(self.output_file)
        print(f"格式化完成！输出文件: {self.output_file}")


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description='基于GB/T 9704-2012标准的Word文档格式化工具')
    parser.add_argument('input_file', help='输入Word文档路径')
    parser.add_argument('output_file', help='输出Word文档路径')
    parser.add_argument('--red-head', action='store_true', help='添加红头标识')
    parser.add_argument('--org', default='国央企', help='单位名称（红头使用）')
    parser.add_argument('--title', default='', help='文件标题（红头使用）')
    parser.add_argument('--number', default='', help='文件编号（红头使用，左对齐）')
    parser.add_argument('--date', default='', help='日期（红头使用）')

    args = parser.parse_args()

    input_file = Path(args.input_file)
    output_file = Path(args.output_file)

    if not input_file.exists():
        print(f"错误: 输入文件不存在 - {input_file}")
        sys.exit(1)

    try:
        formatter = WordFormatter(
            input_file=input_file,
            output_file=output_file,
            org_name=args.org,
            doc_title=args.title,
            doc_number=args.number,
            doc_date=args.date
        )
        formatter.format(add_red_head=args.red_head)
        sys.exit(0)
    except Exception as e:
        print(f"错误: 格式化失败 - {e}")
        sys.exit(1)
